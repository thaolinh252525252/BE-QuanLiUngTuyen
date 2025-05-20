import imaplib
import email
from email.header import decode_header
from email.utils import parseaddr
import os
import re
import unicodedata
from docx import Document
import logging
import google.generativeai as genai
from datetime import datetime
import time
import pdfplumber
import pytesseract
import pdf2image
from PIL import Image
import json
from bson import ObjectId
from pymongo import MongoClient
from services.email_tracker_mongo import init_tracking_collection, has_processed_email, mark_email_as_processed
from services.email_utils import send_result_email
from email.utils import parsedate_to_datetime

EMAIL = "thaovuong669@gmail.com"
APP_PASSWORD = "wgam okla ffjk wxwy"
IMAP_SERVER = "imap.gmail.com"
ATTACHMENTS_DIR = "cv_files"
genai.configure(api_key="AIzaSyAVZplOCSPwJpdtnSeHfPDNstBze_gUZ2Y")

os.makedirs(ATTACHMENTS_DIR, exist_ok=True)
logging.getLogger("pdfplumber").setLevel(logging.ERROR)

POSITIVE_KEYWORDS = [
    "ứng tuyển", "apply", "job application", "cv", "resume", "xin việc", "application", "tuyển dụng", "hồ sơ", "việc làm", "job"
]
NEGATIVE_KEYWORDS = [
    "promotion", "advertisement", "newsletter", "quảng cáo", "khuyến mãi", "sale", "spam",
    "verify", "confirmation", "password", "welcome", "no-reply", "info@", "support@", "notifications@"
]
SPAM_EMAILS = [
    "recommendations@inspire.pinterest.com",
    "no-reply@mail.englishscore.com"
]

def connect_to_mongodb():
    client = MongoClient("${import.meta.env.Mongo_connect}")
    db = client["tuyendung"]
    return db

def safe_filename(filename):
    filename = unicodedata.normalize('NFKD', filename).encode('ascii', 'ignore').decode('utf-8')
    filename = re.sub(r'[^a-zA-Z0-9_.-]', '_', filename)
    return filename

def extract_and_normalize_position(subject):
    if not subject:
        return "Vị trí chưa xác định"

    # Chuẩn hóa văn bản: loại bỏ dấu, chuyển thường, thay khoảng trắng thừa
    subject = unicodedata.normalize("NFKC", subject).strip().lower()
    subject = re.sub(r"\s+", " ", subject)

    # Loại bỏ các cụm từ không cần thiết
    subject = re.sub(r"(?i)(ứng tuyển vị trí|apply for position|vị trí|apply|ứng tuyển)\s*[:]*\s*", "", subject)

    # Tách tên ứng viên và các phần thừa
    patterns = [
        r"\s*[-–]\s*.+$",           # Xóa phần sau dấu "-"
        r"\s*(của|by)\s*.+$",       # Xóa phần sau "của" hoặc "by"
        r"\s+[a-zA-Z\s]+$",         # Xóa tên ở cuối (giả sử tên là chuỗi chữ)
    ]
    for pattern in patterns:
        subject = re.sub(pattern, "", subject)

    # Loại bỏ khoảng trắng thừa và ký tự không mong muốn
    subject = re.sub(r"[^\w\s/]", "", subject).strip()

    # Ánh xạ vị trí sang tên chuẩn
    mappings = {
        r"frontend|front-end|react|vue|ui|giao diện|web.*developer": "Lập trình viên Frontend",
        r"python|flask|django|python|[^a-zA-Z]python[^a-zA-Z]": "Lập trình viên Python",
        r"full[- ]?stack|fullstack|backend.*frontend": "Lập trình viên Full-stack",
        r"sale|bán hàng|kinh doanh|chăm sóc khách|tư vấn": "Nhân viên Kinh doanh",
        r"test|qa|quality assurance|kiểm thử|selenium|junit|testng": "Chuyên viên Kiểm thử phần mềm",
        r"data analyst|phân tích dữ liệu|power bi|tableau|looker|google data studio": "Chuyên viên Phân tích Dữ liệu",
        r"an toàn thông tin|bảo mật|security|cybersecurity|ids|ips|firewall": "Kỹ sư An toàn Thông tin",
        r"devops|ansible|jenkins|terraform|prometheus|grafana|cicd|ci/cd": "Kỹ sư DevOps",
        r"data engineer|big data|spark|hadoop|etl|pipelines|xử lý dữ liệu": "Kỹ sư Dữ liệu",
        r"network|mạng|infrastructure|vpn|lan|wan|router|switch": "Kỹ sư Hạ tầng Mạng",
        r"machine learning|ml|học máy|tensorflow|pytorch": "Kỹ sư Machine Learning",
        r"ai|trí tuệ nhân tạo|artificial intelligence": "Kỹ sư Trí tuệ Nhân tạo",
        r"ui[/-]ux|ux[/-]ui|user experience|user interface|thiết kế giao diện": "Nhà thiết kế UI/UX",
        r"senior.*ui[/-]ux|lead.*ui[/-]ux": "Nhà thiết kế UI/UX Cấp cao"
    }

    for pattern, normalized in mappings.items():
        if re.search(pattern, subject, re.IGNORECASE):
            return normalized

    return subject.title()

def select_matching_jd(vi_tri_ung_tuyen, text, collection=None):
    if collection is None:
        db = connect_to_mongodb()
        collection = db["mo_ta_cong_viec"]

    jds = list(collection.find({}))
    if not jds:
        print("⚠️ Không có JD nào trong database")
        return None

    vi_tri_normalized = vi_tri_ung_tuyen.lower().strip()

    # Ưu tiên exact match
    for jd in jds:
        jd_position = jd.get("vi_tri", "").lower().strip()
        if vi_tri_normalized == jd_position:
            print(f"✅ Ánh xạ chính xác '{vi_tri_ung_tuyen}' → JD: {jd['vi_tri']}")
            return jd

    # Nếu không có exact match, thử khớp gần đúng
    print(f"⚠️ Không có JD exact match cho: {vi_tri_ung_tuyen}, thử khớp gần đúng...")
    best_jd, best_score = None, 0
    tech_keywords = [
        "reactjs", "redux", "tailwindcss", "css", "html", "javascript", "typescript",
        "node.js", "python", "mongodb", "postgresql", "mysql", "graphql", "rest", "api",
        "nessus", "owasp", "siem", "firewall", "ids", "ips", "wireshark", "metasploit",
        "ceh", "cissp", "ci/cd", "jwt", "docker", "kubernetes", "git"
    ]

    for jd_doc in jds:
        jd_position = jd_doc.get("vi_tri", "").lower().strip()
        jd_requirements = jd_doc.get("yeu_cau", "").lower()

        position_score = 0
        if vi_tri_normalized in jd_position or jd_position in vi_tri_normalized:
            position_score = 40
        elif any(keyword in vi_tri_normalized for keyword in jd_position.split()):
            position_score = 20

        jd_skills = set()
        for skill in tech_keywords:
            if skill in jd_requirements:
                jd_skills.add(skill)

        cv_skills = set()
        for word in text.lower().split():
            if word in tech_keywords:
                cv_skills.add(word)

        common_skills = cv_skills.intersection(jd_skills)
        skill_score = len(common_skills) / max(len(jd_skills), 1) * 50

        total_score = position_score + skill_score

        if total_score > best_score:
            best_jd = jd_doc
            best_score = total_score

    if best_jd and best_score >= 60:  # Tăng ngưỡng để đảm bảo độ phù hợp
        print(f"✅ Ánh xạ '{vi_tri_ung_tuyen}' sang JD: {best_jd['vi_tri']} (score: {best_score:.1f}%)")
        return best_jd

    print(f"❌ Không tìm thấy JD phù hợp cho vị trí: {vi_tri_ung_tuyen} (score: {best_score:.1f}%)")
    return None
def safe_get(info, key):
    return (info.get(key) or "").strip()

def pdf_to_text(pdf_path):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    page_text = unicodedata.normalize("NFKC", page_text)
                    page_text = re.sub(r"[^\x00-\x7F\u00C0-\u1EF9\n\s]", " ", page_text)
                    text += page_text + "\n"
            if text.strip():
                return text.strip()
    except Exception as e:
        print(f"❌ Lỗi pdfplumber: {e}")

    try:
        print("📸 Đang chuyển PDF sang ảnh để OCR...")
        images = pdf2image.convert_from_path(pdf_path)
        text = ""
        for image in images:
            ocr_result = pytesseract.image_to_string(image, lang='vie+eng')
            ocr_result = unicodedata.normalize("NFKC", ocr_result)
            ocr_result = re.sub(r"[^\x00-\x7F\u00C0-\u1EF9\n\s]", " ", ocr_result)
            text += ocr_result + "\n"
        return text.strip()
    except Exception as e:
        print(f"❌ Lỗi OCR fallback: {e}")
        return ""

def docx_to_text(docx_path):
    try:
        doc = Document(docx_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_parts.append(cell_text)
        return "\n".join(text_parts).strip()
    except Exception as e:
        print(f"❌ Lỗi đọc DOCX {docx_path}: {e}")
        return ""

def is_recruitment_email(subject: str, from_email: str):
    subject = (subject or "").lower()
    sender = (from_email or "").lower()

    if sender in SPAM_EMAILS or any(bad in sender for bad in NEGATIVE_KEYWORDS):
        return False
    return (
        any(keyword in subject for keyword in POSITIVE_KEYWORDS)
        and not any(keyword in subject for keyword in NEGATIVE_KEYWORDS)
    )

def decode_filename(raw_filename):
    if not raw_filename:
        return None
    decoded_parts = decode_header(raw_filename)
    filename = ""
    for part, encoding in decoded_parts:
        if isinstance(part, bytes):
            filename += part.decode(encoding or "utf-8", errors="ignore")
        else:
            filename += part
    return filename

def normalize_text(text):
    if not text:
        return ""
    # Chỉ làm sạch nhẹ, giữ dấu và từ khóa
    text = text.strip()
    text = re.sub(r"\s+", " ", text).lower()
    return text
def extract_vi_tri_ung_tuyen(subject, text, jd_collection):
    jd_list = list(jd_collection.find({}, {"vi_tri": 1}))
    jd_positions = [normalize_text(jd["vi_tri"]) for jd in jd_list]
    jd_lookup = {normalize_text(jd["vi_tri"]): jd for jd in jd_list}

    match = re.search(r"(?i)ứng tuyển vị trí[:\s]+([^-–\n]+)", subject)
    if match:
        subject_vi_tri = normalize_text(match.group(1))
        print(f"🔎 Vị trí từ subject: {subject_vi_tri}")

        if subject_vi_tri in normalize_text(text):
            print(f"✅ Subject xuất hiện trong CV: {subject_vi_tri}")
            if subject_vi_tri in jd_lookup:
                return subject_vi_tri
            else:
                print("⚠️ Subject xuất hiện trong CV nhưng không có JD khớp")
                return subject_vi_tri

    for jd_pos in jd_positions:
        if jd_pos in normalize_text(text):
            print(f"✅ Tìm thấy JD trong nội dung CV: {jd_pos}")
            return jd_pos

    for jd_pos in jd_positions:
        if jd_pos in normalize_text(subject):
            print(f"⚠️ Subject gần giống JD: {jd_pos}")
            return jd_pos

    print("❌ Không tìm thấy JD chính xác, để Gemini tự đoán")
    return None
def extract_and_normalize_position(subject):
    if not subject:
        return "Vị trí chưa xác định"

    # Chỉ làm sạch nhẹ: giữ nguyên dấu và từ khóa, chỉ xóa tiền tố
    subject = subject.strip()
    
    # Xóa tiền tố như "ứng tuyển vị trí", "apply for position"
    prefixes = r"(?i)(ứng tuyển vị trí|apply for position|vị trí|apply|ứng tuyển)\s*[:]*\s*"
    subject = re.sub(prefixes, "", subject, flags=re.IGNORECASE)

    # Xóa tên ứng viên ở cuối (nếu có dấu gạch ngang hoặc "của/by")
    name_pattern = r"\s*[-–]\s*[a-zA-Z\sàáãạảăắằẳẵặâấầẩẫậèéẹẻẽêềếểễệìíĩỉịòóõọỏôốồổỗộơớờởỡợùúũụủưứừửữựỳýỹỷỵ]+$"  # Chỉ khớp tên tiếng Việt
    subject = re.sub(name_pattern, "", subject)
    
    # Loại bỏ khoảng trắng thừa
    subject = re.sub(r"\s+", " ", subject).strip()

    # Nếu tiêu đề quá ngắn hoặc generic, trả về fallback
    if len(subject.split()) <= 1 or subject.lower() in ["kỹ sư", "ky su"]:
        return "Vị trí chưa xác định"

    # Ánh xạ trực tiếp với JD
    mappings = {
        r"frontend|front-end|react|vue|ui|giao diện|web.*developer": "Lập trình viên Front-end",
        r"python|flask|django|fastapi": "Lập trình viên Python",
        r"full[- ]?stack|fullstack|backend.*frontend": "Lập trình viên Full-stack",
        r"sale|bán hàng|kinh doanh|chăm sóc khách|tư vấn": "Nhân viên Kinh doanh",
        r"test|qa|quality assurance|kiểm thử|selenium|junit|testng": "Chuyên viên Kiểm thử phần mềm",
        r"data analyst|phân tích dữ liệu|power bi|tableau": "Chuyên viên Phân tích Dữ liệu",
        r"an toàn thông tin|bảo mật|security|cybersecurity|ids|ips|firewall": "Kỹ sư An toàn Thông tin",
        r"devops|ansible|jenkins|terraform|prometheus|grafana|cicd|ci/cd": "Kỹ sư DevOps",
        r"data engineer|big data|spark|hadoop|etl|pipelines|xử lý dữ liệu": "Kỹ sư Dữ liệu",
        r"network|mạng|infrastructure|vpn|lan|wan|router|switch": "Kỹ sư Hạ tầng Mạng",
        r"machine learning|ml|học máy|tensorflow|pytorch": "Kỹ sư Machine Learning",
        r"ai|trí tuệ nhân tạo|artificial intelligence": "Kỹ sư Trí tuệ Nhân tạo"
    }

    for pattern, normalized in mappings.items():
        if re.search(pattern, subject, re.IGNORECASE):
            return normalized

    # Nếu không khớp với mappings, trả về tiêu đề gốc đã làm sạch, viết hoa đầu từ
    return subject.title()
def extract_position_and_name_from_subject(subject: str):
    subject = normalize_text(subject)

    pos_match = re.search(r"(?:ung tuyen.*?vi tri|apply for position)\s+(.+?)(?:[-–]|cua|by)", subject)
    position = pos_match.group(1).strip() if pos_match else ""

    name_match = re.search(r"(?:[-–]|cua|by)\s+(.+)$", subject)
    name = name_match.group(1).strip().title() if name_match else ""

    return position, name

def extract_info_with_gemini(text, filename, subject, jd):
    if not jd:
        print("⚠️ Không có JD phù hợp — để Gemini tự phân tích vị trí")
        jd = {
            "vi_tri": "Không xác định",
            "mo_ta": "Không có mô tả cụ thể",
            "yeu_cau": "Không rõ yêu cầu"
        }
        fallback_mode = True
    else:
        fallback_mode = False

    def safe_string(value):
        if not isinstance(value, str):
            value = str(value)
        return value.replace("{", "{{").replace("}", "}}").strip()

    vi_tri = safe_string(jd.get("vi_tri", ""))
    mo_ta = safe_string(jd.get("mo_ta", ""))
    yeu_cau = safe_string(jd.get("yeu_cau", ""))

    jd_section = f"""
- Kiểm tra xem vị trí ứng tuyển (sau khi chuẩn hóa) có khớp với JD ('{vi_tri.lower()}').
- So sánh kỹ năng, kinh nghiệm, chứng chỉ, giải thưởng, dự án, trình độ học vấn với JD:

[Mô tả công việc]
- Vị trí: {vi_tri}
- Mô tả: {mo_ta}
- Yêu cầu: {yeu_cau}

**Hướng dẫn so sánh và tính điểm**:
1. **Kỹ năng (40%)**: Đếm số kỹ năng trong CV khớp với `yeu_cau`. Tính tỷ lệ kỹ năng khớp (số kỹ năng khớp / tổng kỹ năng yêu cầu) và nhân với 40.
2. **Kinh nghiệm (30%)**: Đánh giá số năm kinh nghiệm hoặc số dự án liên quan. Nếu kinh nghiệm >= yêu cầu, cho 30 điểm; nếu đạt 50–99% yêu cầu, cho 15–29 điểm; nếu <50%, cho 0–14 điểm.
3. **Học vấn và chứng chỉ (20%)**: Nếu trình độ học vấn hoặc chứng chỉ khớp với yêu cầu, cho 20 điểm; nếu thiếu một phần, cho 0–19 điểm.
4. **Khác (10%)**: Dự án, giải thưởng hoặc các yếu tố khác liên quan đến JD, tối đa 10 điểm.
5. Tổng điểm (`diem_phu_hop`) từ 0–100
6. Nếu vị trí ứng tuyển không khớp với JD, trả về `nhan_xet`: "Hiện không có JD phù hợp với vị trí này" và `diem_phu_hop` tối đa 50.
"""

    prompt = f"""
Bạn là AI chuyên gia tuyển dụng. Dựa trên CV và JD, hãy trích xuất thông tin theo mẫu JSON sau và đánh giá mức độ phù hợp.

**Văn bản CV**:
{text}

**Trả về JSON chuẩn sau**:
{{
  "ho_ten": "",
  "so_dien_thoai": "",
  "ngay_sinh": "",
  "que_quan": "",
  "noi_o": "",
  "trinh_do_hoc_van": [],
  "kinh_nghiem": [],
  "vi_tri_ung_tuyen": "",
  "ky_nang": [],
  "chung_chi": [],
  "giai_thuong": [],
  "du_an": [],
  "trang_thai": "",
  "diem_phu_hop": 0,
  "nhan_xet": "",
  "muc_tieu_nghe_nghiep": "",
  "so_thich": [],
  "nguoi_gioi_thieu": [],
  "hoat_dong": []
}}

    **Yêu cầu định dạng**:
    - Các trường `trinh_do_hoc_van`, `kinh_nghiem`, `ky_nang`, `chung_chi`, `giai_thuong`, `du_an`, `so_thich`, `nguoi_gioi_thieu`, `hoat_dong` phải là danh sách các CHUỖI (string), KHÔNG phải object.
    - **Chi tiết yêu cầu**:
      - `trinh_do_hoc_van`: Lấy bằng cấp, chuyên ngành, trường và năm tốt nghiệp (nếu có). Ví dụ: ["Cử nhân CNTT Đại học Bách khoa Hà Nội 2010"].
      - `kinh_nghiem`: Lấy vị trí công việc, công ty, thời gian và mô tả ngắn gọn nhiệm vụ. Ví dụ: ["Network Engineer tại Viettel 2018-2020: Cấu hình router Cisco"].
      - `ky_nang`: Lấy các kỹ năng kỹ thuật hoặc mềm. Ví dụ: ["Python", "Teamwork"].
      - `chung_chi`: Lấy tên chứng chỉ, năm và tổ chức cấp (nếu có). Ví dụ: ["TOEIC 900 2023", "AWS Certified 2022"].
      - `giai_thuong`: Lấy tên giải thưởng, năm và tổ chức/sự kiện. Ví dụ: ["Giải nhất Hackathon 2023"].
      - `du_an`: Lấy tên dự án, công ty, năm và mô tả ngắn gọn nhiệm vụ, gộp thành một chuỗi. Ví dụ: ["Xây dựng mạng nội bộ cho trung tâm dữ liệu tại IDC Việt Nam 2023: Cấu hình BGP, OSPF, Firewall đa lớp"].
      - `so_thich`: Lấy sở thích cá nhân. Ví dụ: ["Tập gym", "Tham gia hackathon"].
      - `nguoi_gioi_thieu`: Lấy tên, vị trí, công ty, email và số điện thoại (nếu có). Ví dụ: ["Nguyễn Văn A - Manager tại Viettel - a@viettel.vn - 0123456789"].
      - `hoat_dong`: Lấy vai trò, tổ chức, thời gian và mô tả ngắn gọn. Ví dụ: ["Trưởng ban Câu lạc bộ Khởi nghiệp 2018-2020: Tổ chức workshop"].
    - Nếu không tìm thấy thông tin, để danh sách rỗng (`[]`) hoặc chuỗi rỗng (`""`).


**Lưu ý**:
- Chỉ trả về đúng JSON, không thêm mô tả ngoài.
- `diem_phu_hop` phải từ 0–100, tính theo hướng dẫn trên.
- `nhan_xet` phải nêu rõ những điểm CV còn thiếu so với JD và gợi ý cải thiện. Nếu vị trí không khớp, ghi "Hiện không có JD phù hợp với vị trí này".
- Không điền thông tin nếu không có trong CV.
- Các trường mới như `muc_tieu_nghe_nghiep`, `so_thich`, `nguoi_gioi_thieu`, `hoat_dong` chỉ điền nếu có dữ liệu trong CV.
- ⚠️ Lưu ý: Vì không tìm thấy JD phù hợp nên bạn phải tự xác định: đầu tiên, xác định trên tiêu đề (tự đọc để nhận biết, có thể bỏ từ không liên qua như vị trí, ứng tuyển, vị trí ứng tuyển hoặc của, Tên người gửi). 
Nếu lấy được vị trí ứng tuyển ở tiêu đề mới đọc từ nội dung CV, xem trong 10 dòng đầu có câu nào tương tự với tiêu đề ko, nếu có đó là vị trí ứng tuyển, nếu không tự suy ra từ nội dung CV (kiểm tra tầm 10 dòng đầu để suy ra thôi). Nếu không xác định được, hãy để trống trường `vi_tri_ung_tuyen`. Trong mọi trường hợp, `nhan_xet` phải ghi rõ 'Không tìm thấy JD phù hợp' và điểm tối đa không vượt quá 50.
- Nếu `ho_ten` = null . Đầu tiên bạn xem lại tiêu đề (thường thì tên sau chữ của hoăc -), nếu có thì gán đó là `ho_ten`, nếu không có xem xét ở nội dung CV tự đọc xem trong vòng 10 dòng đầu tiên có phần nào nổi bật, giống họ tên người nhất thì gán, không để `ho_ten` = null


{jd_section}
"""

    try:
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(prompt)
        print("📜 Phản hồi thô từ Gemini:", response.text)

        text = response.text.strip()
        if "```" in text:
            parts = [part for part in text.split("```") if "{" in part]
            raw_json = parts[0] if parts else text
        else:
            raw_json = text

        match = re.search(r"\{.*\}", raw_json, re.DOTALL)
        if not match:
            print("❌ Không tìm thấy JSON hợp lệ trong phản hồi Gemini")
            return None

        result = json.loads(match.group())

        if "diem_phu_hop" not in result or not isinstance(result["diem_phu_hop"], (int, float)):
            print("❌ Gemini không trả về trường 'diem_phu_hop' hợp lệ, trả về giá trị mặc định")
            result["diem_phu_hop"] = 0
            result["nhan_xet"] = "Không thể đánh giá mức độ phù hợp do thiếu thông tin từ Gemini."

        return result

    except Exception as e:
        print(f"❌ Lỗi khi gọi Gemini: {e}")
        import traceback
        traceback.print_exc()
        return None

def process_all_emails(mongo_collection):
    tracking_collection = init_tracking_collection(mongo_collection.database)
    from_email = "Không có email được xử lý"
    try:
        mail = imaplib.IMAP4_SSL(IMAP_SERVER)
        mail.login(EMAIL, APP_PASSWORD)
        mail.select("inbox")

        status, messages = mail.search(None, "ALL")
        if status != "OK":
            print("❌ Không tìm thấy email.")
            return
        print("🟢 Bắt đầu quét email...")

        for num in messages[0].split():
            _, msg_data = mail.fetch(num, "(RFC822)")
            msg = email.message_from_bytes(msg_data[0][1])
            uid = msg["Message-ID"] or num.decode()
            
            email_date_raw = msg["Date"]
            try:
                email_date = parsedate_to_datetime(email_date_raw)
            except:
                email_date = datetime.now()

            print("📌 UID của email:", uid)
            print("🔍 Tìm trong tracking:", tracking_collection.find_one({"uid": uid}))

            subject = decode_header(msg.get("Subject") or "")[0][0]
            if isinstance(subject, bytes):
                subject = subject.decode(errors="ignore")
            from_email = parseaddr(msg.get("From"))[1]

            print(f"📨 Email: {subject} -- From: {from_email}")
            print("🔎 is_recruitment_email:", is_recruitment_email(subject, from_email))
            print("🔄 has_processed_email:", has_processed_email(tracking_collection, uid))

            if not is_recruitment_email(subject, from_email) or has_processed_email(tracking_collection, uid):
                print("⛔ Email bị bỏ qua")
                continue

            text, filepath = "", None

            for part in msg.walk():
                if part.get_content_maintype() == "multipart":
                    continue
                if part.get("Content-Disposition") is None:
                    continue

                filename = decode_filename(part.get_filename())
                if filename and filename.lower().endswith((".pdf", ".docx")):
                    safe_name = safe_filename(filename)
                    filepath = os.path.join(ATTACHMENTS_DIR, safe_name)
                    with open(filepath, "wb") as f:
                        f.write(part.get_payload(decode=True))
                    if filename.endswith(".pdf"):
                        text = pdf_to_text(filepath)
                    else:
                        text = docx_to_text(filepath)

            if not text:
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        charset = part.get_content_charset() or "utf-8"
                        try:
                            text += part.get_payload(decode=True).decode(charset, errors="ignore")
                        except:
                            pass

            jd_collection = mongo_collection.database["mo_ta_cong_viec"]
            jd_docs = list(jd_collection.find({}, {"vi_tri": 1}))
            jd_positions = [normalize_text(doc["vi_tri"]) for doc in jd_docs]

           # Lấy vị trí từ subject
            vi_tri_from_subject = extract_and_normalize_position(subject)
            print(f"🧠 Vị trí chuẩn hóa từ subject: {vi_tri_from_subject}")

            # Kiểm tra vị trí trong nội dung CV (10 dòng đầu)
            vi_tri_from_text = ""
            lines = text.strip().split("\n")
            jd_collection = mongo_collection.database["mo_ta_cong_viec"]
            jd_positions = [jd["vi_tri"].lower().strip() for jd in jd_collection.find({}, {"vi_tri": 1})]

            for line in lines[:10]:
                line_norm = line.lower().strip()
                for jd_pos in jd_positions:
                    if jd_pos in line_norm:
                        vi_tri_from_text = jd_pos
                        break
                if vi_tri_from_text:
                    break

            # Chọn vị trí cuối cùng
            vi_tri_chuan_norm = vi_tri_from_subject if vi_tri_from_subject != "Vị trí chưa xác định" else vi_tri_from_text
            if not vi_tri_chuan_norm:
                vi_tri_chuan_norm = subject.strip().title()  # Fallback to cleaned subject
            print(f"🧠 Vị trí chuẩn hóa cuối cùng: {vi_tri_chuan_norm}")

            # Tìm JD phù hợp
            matched_jd = select_matching_jd(vi_tri_chuan_norm, text, jd_collection)

            info = extract_info_with_gemini(
                text,
                os.path.basename(filepath) if filepath else "no_file",
                subject,
                matched_jd
            )

            trang_thai_ung_vien = ""
            status_match = re.search(r"(?i)trạng thái.*?:\s*([^\n]+)", text)
            if status_match:
                trang_thai_ung_vien = status_match.group(1).strip()
            elif "chờ phỏng vấn" in text.lower():
                trang_thai_ung_vien = "Đang ứng tuyển"
            subject_vi_tri, subject_name = extract_position_and_name_from_subject(subject)
            if not info or not info.get("ho_ten"):
                if subject_name:
                    info = info or {}
                    info["ho_ten"] = subject_name
                    print(f"📝 Tên lấy từ tiêu đề: {info['ho_ten']}")
                else:
                    print("⚠️ Không tìm thấy tên ứng viên, bỏ qua email")
                    continue

            if not info.get("vi_tri_ung_tuyen") and subject_vi_tri:
                info["vi_tri_ung_tuyen"] = subject_vi_tri.title()
                print(f"📌 Vị trí ứng tuyển fallback từ subject: {info['vi_tri_ung_tuyen']}")

            diem_phu_hop = info.get("diem_phu_hop")
            doc = {
            "ho_ten": info.get("ho_ten", ""),
            "email": from_email,
            "so_dien_thoai": safe_get(info, "so_dien_thoai"),
            "ngay_sinh": safe_get(info, "ngay_sinh"),
            "que_quan": safe_get(info, "que_quan"),
            "noi_o": safe_get(info, "noi_o"),
            "trinh_do_hoc_van": info.get("trinh_do_hoc_van", []),
            "kinh_nghiem": info.get("kinh_nghiem", []),
            "vi_tri_ung_tuyen": matched_jd["vi_tri"] if matched_jd else vi_tri_chuan_norm,
            "ky_nang": info.get("ky_nang", []),
            "chung_chi": info.get("chung_chi", []),
            "giai_thuong": info.get("giai_thuong", []),
            "du_an": info.get("du_an", []),
            "trang_thai": trang_thai_ung_vien or "Đang ứng tuyển",
            "trang_thai_gui_email": "Chưa gửi",
            "ngay_nop": datetime.now(),
            "diem_phu_hop": int(diem_phu_hop or 0),
            "nhan_xet": info.get("nhan_xet", "Hiện không có JD phù hợp với vị trí này. Đề nghị xem xét các vị trí như Kỹ sư Trí tuệ Nhân tạo hoặc Kỹ sư Machine Learning." if not matched_jd else ""),
            "cv_filepath": filepath or "",
            "ngay_gui": email_date,
            "muc_tieu_nghe_nghiep": info.get("muc_tieu_nghe_nghiep", ""),
            "so_thich": info.get("so_thich", []),
            "nguoi_gioi_thieu": info.get("nguoi_gioi_thieu", []),
            "hoat_dong": info.get("hoat_dong", [])
        }

            mongo_collection.insert_one(doc)
            mark_email_as_processed(tracking_collection, uid, from_email, subject)
            print(f"✅ Đã lưu ứng viên: {doc['ho_ten']} - {from_email} - Điểm: {diem_phu_hop}")
            time.sleep(5)

        mail.logout()

    except Exception as e:
        print(f"❌ Lỗi xử lý email: {e}")
        import traceback
        traceback.print_exc()

    print(f"✅ Đã xử lý email từ: {from_email}")