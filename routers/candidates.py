from pymongo import MongoClient
from bson import ObjectId, errors
from services.email_utils import send_result_email, send_interview_email
from fastapi import APIRouter, HTTPException, Query, Body, Form, UploadFile, File
from datetime import datetime
import os
import unicodedata
import re
import pdfplumber
import pytesseract
import pdf2image
import google.generativeai as genai
import json
from fastapi.responses import JSONResponse
from collections import defaultdict
from datetime import datetime, timedelta
from docx import Document

router = APIRouter()

# Cấu hình MongoDB
try:
    client = MongoClient("${import.meta.env.Mongo_connect}", serverSelectionTimeoutMS=3000)
    db = client["tuyendung"]
    collection = db["ung_vien"]
    jd_collection = db["mo_ta_cong_viec"]
    client.server_info()
    print("✅ Đã kết nối MongoDB")
except Exception as e:
    print("❌ Lỗi kết nối MongoDB:", e)

# Cấu hình thư mục lưu file và Gemini
UPLOAD_DIR = "cv_files"
os.makedirs(UPLOAD_DIR, exist_ok=True)
genai.configure(api_key="AIzaSyAVZplOCSPwJpdtnSeHfPDNstBze_gUZ2Y")  # Thay bằng API key của bạn

# Hàm chuẩn hóa tên file
def safe_filename(filename):
    filename = unicodedata.normalize('NFKD', filename).encode('ascii', 'ignore').decode('utf-8')
    filename = re.sub(r'[^\w\-_\. ]', '_', filename)
    name, ext = os.path.splitext(filename)
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    final_name = f"{name}_{timestamp}{ext}"  # Tên lưu trữ
    original_name = f"{name}{ext}"  # Tên gốc
    return original_name, final_name  # Trả về cả hai

# Hàm đọc nội dung PDF
def pdf_to_text(pdf_path):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    page_text = unicodedata.normalize("NFKC", page_text)
                    page_text = re.sub(r'\n\s*\n', '\n', page_text)  # Loại bỏ dòng trống thừa
                    page_text = re.sub(r"[^\x00-\x7F\u00C0-\u1EF9\n\s]", " ", page_text)
                    text += page_text + "\n"
            if text.strip():
                print("📄 Nội dung PDF trích xuất:", text)
                return text.strip()
    except Exception as e:
        print(f"❌ Lỗi pdfplumber: {e}")

    try:
        print("📸 Đang chuyển PDF sang ảnh để OCR...")
        images = pdf2image.convert_from_path(pdf_path)
        text = ""
        for image in images:
            ocr_result = pytesseract.image_to_string(image, lang='vie+eng')
            ocr_result = unicodedata.normalize("NFKC", ocr_result)
            ocr_result = re.sub(r"[^\x00-\x7F\u00C0-\u1EF9\n\s]", " ", ocr_result)
            text += ocr_result + "\n"
        print("📄 Nội dung OCR:", text)
        return text.strip()
    except Exception as e:
        print(f"❌ Lỗi OCR fallback: {e}")
        return ""

# Hàm đọc nội dung DOCX
def docx_to_text(docx_path):
    try:
        doc = Document(docx_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_parts.append(cell_text)
        return "\n".join(text_parts).strip()
    except Exception as e:
        print(f"❌ Lỗi đọc DOCX {docx_path}: {e}")
        return ""

# Hàm tìm JD phù hợp
def select_matching_jd(position):
    try:
        jd = jd_collection.find_one({"vi_tri": position})
        if jd:
            print(f"📌 Tìm thấy JD khớp chính xác: {jd['vi_tri']}")
            return jd

        jds = jd_collection.find({"vi_tri": {"$regex": re.escape(position), "$options": "i"}})
        best_jd = None
        best_score = 0
        position_lower = position.lower()

        for jd in jds:
            jd_position = jd.get("vi_tri", "").lower()
            common_length = len(set(position_lower.split()) & set(jd_position.split()))
            score = common_length / max(len(position_lower.split()), 1)
            if score > best_score:
                best_score = score
                best_jd = jd

        if best_jd and best_score > 0.5:
            print(f"📌 Tìm thấy JD gần giống: {best_jd['vi_tri']} (điểm khớp: {best_score})")
            return best_jd
        else:
            print(f"⚠️ Không tìm thấy JD phù hợp cho vị trí: {position}")
            return None
    except Exception as e:
        print(f"❌ Lỗi khi tìm JD: {e}")
        return None

# Hàm trích xuất thông tin bổ sung với Gemini
def extract_info_with_gemini(text, filename, position, jd=None):
    jd_text = ""
    if jd:
        jd_text = f"""
        **Mô tả công việc (JD)**:
        Vị trí: {jd.get('vi_tri', '')}
        Yêu cầu: {jd.get('mo_ta', '')}
        Kỹ năng yêu cầu: {', '.join(jd.get('ky_nang', []))}
        Chứng chỉ ưu tiên: {', '.join(jd.get('chung_chi', []))}
        """

    prompt = f"""
    Bạn là AI chuyên gia tuyển dụng. Dựa trên CV và mô tả công việc (nếu có), hãy trích xuất các thông tin bổ sung theo mẫu JSON sau. Chỉ trích xuất các trường được yêu cầu, bỏ qua các trường đã được cung cấp từ form (ho_ten, email, so_dien_thoai, vi_tri_ung_tuyen, ngay_sinh, que_quan, noi_o).

    **Mẫu JSON**:
    {{
      "trinh_do_hoc_van": [],
      "kinh_nghiem": [],
      "ky_nang": [],
      "chung_chi": [],
      "giai_thuong": [],
      "du_an": [],
      "muc_tieu_nghe_nghiep": "",
      "so_thich": [],
      "nguoi_gioi_thieu": [],
      "hoat_dong": [],
      "diem_phu_hop": 0,
      "nhan_xet": ""
    }}

    **Yêu cầu định dạng**:
    - Các trường `trinh_do_hoc_van`, `kinh_nghiem`, `ky_nang`, `chung_chi`, `giai_thuong`, `du_an`, `so_thich`, `nguoi_gioi_thieu`, `hoat_dong` phải là danh sách các CHUỖI (string), KHÔNG phải object.
    - **Chi tiết yêu cầu**:
      - `trinh_do_hoc_van`: Lấy bằng cấp, chuyên ngành, trường và năm tốt nghiệp (nếu có). Ví dụ: ["Cử nhân CNTT Đại học Bách khoa Hà Nội 2010"].
      - `kinh_nghiem`: Lấy vị trí công việc, công ty, thời gian và mô tả ngắn gọn nhiệm vụ. Ví dụ: ["Network Engineer tại Viettel 2018-2020: Cấu hình router Cisco"].
      - `ky_nang`: Lấy các kỹ năng kỹ thuật hoặc mềm. Ví dụ: ["Python", "Teamwork"].
      - `chung_chi`: Lấy tên chứng chỉ, năm và tổ chức cấp (nếu có). Ví dụ: ["TOEIC 900 2023", "AWS Certified 2022"].
      - `giai_thuong`: Lấy tên giải thưởng, năm và tổ chức/sự kiện. Ví dụ: ["Giải nhất Hackathon 2023"].
      - `du_an`: Lấy tên dự án, công ty, năm và mô tả ngắn gọn nhiệm vụ, gộp thành một chuỗi. Ví dụ: ["Xây dựng mạng nội bộ cho trung tâm dữ liệu tại IDC Việt Nam 2023: Cấu hình BGP, OSPF, Firewall đa lớp"].
      - `so_thich`: Lấy sở thích cá nhân. Ví dụ: ["Tập gym", "Tham gia hackathon"].
      - `nguoi_gioi_thieu`: Lấy tên, vị trí, công ty, email và số điện thoại (nếu có). Ví dụ: ["Nguyễn Văn A - Manager tại Viettel - a@viettel.vn - 0123456789"].
      - `hoat_dong`: Lấy vai trò, tổ chức, thời gian và mô tả ngắn gọn. Ví dụ: ["Trưởng ban Câu lạc bộ Khởi nghiệp 2018-2020: Tổ chức workshop"].
    - Nếu không tìm thấy thông tin, để danh sách rỗng (`[]`) hoặc chuỗi rỗng (`""`).

    **Văn bản CV**:
    {text}

    {jd_text}

    **Lưu ý**:
    - Chỉ trả về đúng JSON, không thêm mô tả ngoài.
    - `diem_phu_hop` (0–100): Tính dựa trên mức độ phù hợp giữa CV và JD (nếu có). So sánh kỹ năng, kinh nghiệm, chứng chỉ, và dự án trong CV với yêu cầu JD. Nếu không có JD, trả về 0.
    - `nhan_xet`: Nêu điểm mạnh/yếu của CV so với JD (nếu có) và gợi ý cải thiện. Nếu không có JD, chỉ nhận xét chung về CV.
    - Trích xuất đầy đủ và chi tiết tất cả các trường, đặc biệt là `du_an`, gộp thông tin thành chuỗi duy nhất.
    """
    try:
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(prompt)
        raw_json = response.text.strip()
        if "```" in raw_json:
            raw_json = [part for part in raw_json.split("```") if "{" in part][0]
        match = re.search(r"\{.*\}", raw_json, re.DOTALL)
        if not match:
            print("❌ Không tìm thấy JSON hợp lệ")
            return None
        result = json.loads(match.group())

        # Xử lý các danh sách để đảm bảo chỉ chứa chuỗi
        list_fields = [
            "trinh_do_hoc_van", "kinh_nghiem", "ky_nang", "chung_chi",
            "giai_thuong", "du_an", "so_thich", "nguoi_gioi_thieu", "hoat_dong"
        ]
        for field in list_fields:
            if field in result and isinstance(result[field], list):
                result[field] = [
                    str(item) if isinstance(item, str) else f"{item.get('name', '')} {item.get('year', '')} {item.get('score', '')}".strip()
                    for item in result[field]
                ]
                result[field] = [item for item in result[field] if item]

        print("📜 Thông tin trích xuất từ Gemini:", result)
        return result
    except Exception as e:
        print(f"❌ Lỗi khi gọi Gemini: {e}")
        return None

# Endpoint GET danh sách ứng viên
@router.get("/")
def get_candidates():
    try:
        print("🔗 Đang truy vấn dữ liệu ứng viên")
        docs = collection.find({"ho_ten": {"$ne": ""}})
        result = []
        for doc in docs:
            doc["id"] = str(doc["_id"])
            del doc["_id"]
            result.append(doc)
        print("📊 Dữ liệu trả về từ GET /candidates/:", result)
        return result
    except Exception as e:
        print("❌ Lỗi khi truy vấn Mongo:", e)
        return {"error": str(e)}

# Endpoint POST để thêm ứng viên mới
@router.post("/")
async def add_candidate(
    ho_ten: str = Form(...),
    email: str = Form(...),
    so_dien_thoai: str = Form(...),
    vi_tri_ung_tuyen: str = Form(...),
    ngay_sinh: str = Form(None),
    que_quan: str = Form(None),
    noi_o: str = Form(None),
    cv_file: UploadFile = File(None),
):
    try:
        cv_filepath = ""
        info = None

        # Tìm JD phù hợp
        jd = select_matching_jd(vi_tri_ung_tuyen)

        # Xử lý file CV nếu có
        if cv_file:
            file_extension = cv_file.filename.lower().split(".")[-1]
            if file_extension not in ["pdf", "docx", "doc"]:
                raise HTTPException(status_code=400, detail="Chỉ chấp nhận file PDF hoặc Word (DOCX)")
            
            if cv_file.size > 5 * 1024 * 1024:
                raise HTTPException(status_code=400, detail="File vượt quá 5MB")

            # Lấy cả tên gốc và tên lưu trữ
            original_filename, storage_filename = safe_filename(cv_file.filename)
            cv_filepath = os.path.join(UPLOAD_DIR, original_filename).replace("\\", "/")  # Lưu tên gốc
            storage_path = os.path.join(UPLOAD_DIR, storage_filename).replace("\\", "/")  # Đường dẫn lưu trữ thực tế
            with open(storage_path, "wb") as f:
                f.write(await cv_file.read())

            print(f"📂 cv_filepath (tên gốc): {cv_filepath}")  # Log để debug
            print(f"📂 storage_path (tên lưu trữ): {storage_path}")  # Log để debug

            # Trích xuất nội dung từ file
            text = ""
            if file_extension == "pdf":
                text = pdf_to_text(storage_path)
            elif file_extension in ["docx", "doc"]:
                text = docx_to_text(storage_path)
            
            if not text:
                print("⚠️ Không trích xuất được nội dung từ file")

            # Trích xuất thông tin bổ sung với Gemini
            info = extract_info_with_gemini(text, cv_file.filename, vi_tri_ung_tuyen, jd) if text else None

        # Tạo document cho MongoDB
        new_candidate = {
            "ho_ten": ho_ten,
            "email": email,
            "so_dien_thoai": so_dien_thoai,
            "vi_tri_ung_tuyen": vi_tri_ung_tuyen,
            "ngay_sinh": ngay_sinh,
            "que_quan": que_quan,
            "noi_o": noi_o,
            "trinh_do_hoc_van": info.get("trinh_do_hoc_van", []) if info else [],
            "kinh_nghiem": info.get("kinh_nghiem", []) if info else [],
            "ky_nang": info.get("ky_nang", []) if info else [],
            "chung_chi": info.get("chung_chi", []) if info else [],
            "giai_thuong": info.get("giai_thuong", []) if info else [],
            "du_an": info.get("du_an", []) if info else [],
            "trang_thai": "Đang ứng tuyển",
            "ngay_nop": datetime.now(),
            "diem_phu_hop": info.get("diem_phu_hop", 0) if info else 0,
            "nhan_xet": info.get("nhan_xet", "") if info else "",
            "cv_filepath": cv_filepath,
            "muc_tieu_nghe_nghiep": info.get("muc_tieu_nghe_nghiep", "") if info else "",
            "so_thich": info.get("so_thich", []) if info else [],
            "nguoi_gioi_thieu": info.get("nguoi_gioi_thieu", []) if info else [],
            "hoat_dong": info.get("hoat_dong", []) if info else []
        }

        # Lưu vào MongoDB
        result = collection.insert_one(new_candidate)

        # Truy vấn lại để lấy document đã insert
        doc = collection.find_one({"_id": result.inserted_id})
        doc["id"] = str(doc["_id"])
        del doc["_id"]

        # In danh sách dự án
        print("Dự án        :")
        du_an_list = doc.get("du_an", [])

        if not isinstance(du_an_list, list):
            du_an_list = [du_an_list]  # Biến thành list dù ban đầu là chuỗi hoặc None

        for d in du_an_list:
            print(f"  - {d}")

        print("📦 Document lưu vào MongoDB:", doc)

        # Xóa file tạm
        if storage_path and os.path.exists(storage_path):
            os.remove(storage_path)
            print(f"🗑️ Đã xóa file tạm: {storage_path}")

        return doc

    except Exception as e:
        print("❌ Lỗi thêm ứng viên:", e)
        raise HTTPException(status_code=500, detail="Không thể thêm ứng viên")


@router.get("/interviews")
def get_all_interviews():
    try:
        # Thêm log để debug
        print("📋 Đang truy vấn lịch phỏng vấn từ database...")
        
        docs = collection.find({
            "lich_phong_van.ngay": {"$exists": True, "$ne": ""},
            "lich_phong_van.gio": {"$exists": True, "$ne": ""},
            "lich_phong_van.dia_diem": {"$exists": True, "$ne": ""}
        })

        result = []
        for doc in docs:
            lich = doc.get("lich_phong_van", {})
            result.append({
                "id": str(doc["_id"]),
                "candidate": doc.get("ho_ten", "Không tên"),
                "date": lich.get("ngay", ""),
                "time": lich.get("gio", ""),
                "room": lich.get("dia_diem", "")
            })
        
        print(f"✅ Tìm thấy {len(result)} lịch phỏng vấn")
        return result
    except Exception as e:
        print(f"❌ Lỗi khi truy vấn lịch phỏng vấn: {e}")
        # Trả về lỗi với status code 500
        from fastapi import HTTPException
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/potential")
def get_potential_candidates(position: str = Query(...)):
    # Tìm ứng viên có vị trí đúng (không phân biệt hoa thường)
    query = {
        "vi_tri_ung_tuyen": { "$regex": f"^{position}$", "$options": "i" }
    }

    # Sắp xếp theo điểm từ cao xuống
    docs = list(collection.find(query).sort("diem_phu_hop", -1))

    # Convert ObjectId -> string
    for doc in docs:
        doc["id"] = str(doc["_id"])
        del doc["_id"]
    return docs

@router.get("/potential-positions")
def get_valid_candidate_positions():
    jd_positions = db["mo_ta_cong_viec"].distinct("vi_tri")
    jd_positions_lower = [p.lower() for p in jd_positions if isinstance(p, str)]

    # So sánh không phân biệt hoa thường
    candidate_positions = db["ung_vien"].distinct("vi_tri_ung_tuyen")
    valid = [p for p in candidate_positions if isinstance(p, str) and p.lower() in jd_positions_lower]

    return sorted(valid)

@router.post("/schedule-interview")
def schedule_interview(
    data: dict = Body(...)
):
    try:
        print(f"📝 Nhận dữ liệu đặt lịch: {data}")
        
        candidate_ids = data.get("candidateIds", [])
        interview_date = data.get("date")
        interview_time = data.get("time")
        interview_location = data.get("location")
        
        print(f"📌 Thông tin đặt lịch: {len(candidate_ids)} ứng viên, ngày: {interview_date}, giờ: {interview_time}, địa điểm: {interview_location}")
        
        if not candidate_ids or not interview_date or not interview_time or not interview_location:
            print("❌ Thiếu thông tin cần thiết")
            raise HTTPException(status_code=400, detail="Thiếu thông tin cần thiết")
        
        # Tạo đối tượng datetime từ date và time
        try:
            interview_datetime = datetime.strptime(f"{interview_date} {interview_time}", "%Y-%m-%d %H:%M")
            print(f"✅ Thời gian phỏng vấn hợp lệ: {interview_datetime}")
        except ValueError as e:
            print(f"❌ Định dạng ngày giờ không hợp lệ: {e}")
            raise HTTPException(status_code=400, detail=f"Định dạng ngày giờ không hợp lệ: {str(e)}")
        
        # Kiểm tra xem collection có tồn tại không
        print(f"📊 Kiểm tra collection: {collection}")
        
        # Chuyển đổi các ID thành ObjectId
        object_ids = []
        for id_str in candidate_ids:
            try:
                object_id = ObjectId(id_str)
                object_ids.append(object_id)
                print(f"✅ ID hợp lệ: {id_str}")
            except errors.InvalidId:
                print(f"❌ ID không hợp lệ: {id_str}")
                continue
        
        if not object_ids:
            print("❌ Không có ID ứng viên hợp lệ")
            raise HTTPException(status_code=400, detail="Không có ID ứng viên hợp lệ")
        
        # Tìm các ứng viên được chọn
        print(f"🔍 Tìm kiếm {len(object_ids)} ứng viên trong database")
        candidates = list(collection.find({"_id": {"$in": object_ids}}))
        print(f"✅ Tìm thấy {len(candidates)} ứng viên")
        
        # Phân loại ứng viên
        valid_candidates = []
        invalid_candidates = []
        
        for candidate in candidates:
            candidate_name = candidate.get("ho_ten", "Không tên")
            candidate_result = candidate.get("ket_qua", "Chưa có kết quả")
            print(f"👤 Kiểm tra ứng viên {candidate_name}, kết quả CV: {candidate_result}")
            
            # Chỉ đặt lịch cho ứng viên đã qua vòng CV (kết quả là "Pass")
            if candidate.get("ket_qua") == "Pass":
                valid_candidates.append(candidate)
                print(f"✅ Ứng viên {candidate_name} đủ điều kiện")
            else:
                invalid_candidates.append(candidate)
                print(f"❌ Ứng viên {candidate_name} không đủ điều kiện (kết quả: {candidate_result})")
        
        if not valid_candidates:
            print("❌ Không có ứng viên nào đủ điều kiện để đặt lịch phỏng vấn")
            return {
                "success": False,
                "message": "Không có ứng viên nào đủ điều kiện để đặt lịch phỏng vấn. Vui lòng chọn ứng viên đã Pass CV."
            }
        
        valid_ids = [candidate["_id"] for candidate in valid_candidates]
        
        # Cập nhật thông tin lịch phỏng vấn cho các ứng viên hợp lệ
        print(f"📝 Cập nhật lịch phỏng vấn cho {len(valid_candidates)} ứng viên")
        try:
            update_result = collection.update_many(
                {"_id": {"$in": valid_ids}},
                {"$set": {
                    "lich_phong_van": {
                        "ngay": interview_date,
                        "gio": interview_time,
                        "dia_diem": interview_location,
                        "trang_thai": "Đã lên lịch"
                    },
                    "trang_thai": "Chờ phỏng vấn"  # Cập nhật trạng thái ứng viên
                }}
            )
            print(f"✅ Đã cập nhật {update_result.modified_count} ứng viên")
        except Exception as e:
            print(f"❌ Lỗi khi cập nhật database: {str(e)}")
            import traceback
            print(traceback.format_exc())
            raise HTTPException(status_code=500, detail=f"Lỗi khi cập nhật thông tin: {str(e)}")
        
        # Gửi email thông báo lịch phỏng vấn cho ứng viên hợp lệ
        email_success = []
        email_failed = []
        
        for candidate in valid_candidates:
            email = candidate.get("email")
            ho_ten = candidate.get("ho_ten", "Ứng viên")
            
            if not email:
                print(f"⚠️ Ứng viên {ho_ten} không có email")
                email_failed.append(f"{ho_ten} (không có email)")
                continue
                
            try:
                print(f"📧 Gửi email thông báo đến {ho_ten} ({email})")
                send_interview_email(
                    email,
                    ho_ten,
                    interview_date,
                    interview_time,
                    interview_location
                )
                email_success.append(ho_ten)
                print(f"✅ Đã gửi email cho {ho_ten}")
            except Exception as e:
                print(f"❌ Lỗi gửi email cho {ho_ten} ({email}): {str(e)}")
                import traceback
                print(traceback.format_exc())
                email_failed.append(f"{ho_ten} ({str(e)})")
        
        # Chuẩn bị thông báo
        message = f"Đã đặt lịch phỏng vấn cho {len(valid_candidates)} ứng viên."
        
        if email_success:
            message += f"\n\nĐã gửi email thông báo đến {len(email_success)} ứng viên."
        
        if email_failed:
            message += f"\n\nKhông thể gửi email đến {len(email_failed)} ứng viên: {', '.join(email_failed)}"
            
        if invalid_candidates:
            invalid_names = [c.get("ho_ten", "Không tên") for c in invalid_candidates]
            message += f"\n\nCác ứng viên không đủ điều kiện (chưa Pass CV): {', '.join(invalid_names)}"
        
        print(f"✅ Hoàn thành đặt lịch phỏng vấn")
        return {
            "success": True,
            "valid_count": len(valid_candidates),
            "invalid_count": len(invalid_candidates),
            "email_success": len(email_success),
            "email_failed": len(email_failed),
            "message": message
        }
    except HTTPException:
        # Đã xử lý ở trên, chỉ ném lại
        raise
    except Exception as e:
        print(f"❌ Lỗi không xác định: {str(e)}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Lỗi server: {str(e)}")



@router.get("/stats")
def get_candidate_stats():
    stats = {
        "total": collection.count_documents({}),
        "by_status": {},
        "by_position": {},
        "by_month": {},  # Thêm dòng này
        "scheduled_count": collection.count_documents({"lich_phong_van": {"$exists": True}}),
        "high_score_count": collection.count_documents({"diem_phu_hop": {"$gte": 80}})
    }

    # Thống kê theo trạng thái
    status_list = collection.aggregate([
        {"$group": {"_id": "$ket_qua", "count": {"$sum": 1}}}
    ])
    for item in status_list:
        key = item["_id"] if item["_id"] else "Chưa đánh giá"
        stats["by_status"][key] = item["count"]

    # Thống kê theo vị trí ứng tuyển
    position_list = collection.aggregate([
        {"$group": {"_id": "$vi_tri_ung_tuyen", "count": {"$sum": 1}}}
    ])
    for item in position_list:
        key = item["_id"] if item["_id"] else "Không rõ"
        stats["by_position"][key] = item["count"]

    # Thống kê theo tháng từ trường `ngay_nop` (hoặc `ngay_gui`)
    month_counts = defaultdict(int)
    for doc in collection.find({"ngay_nop": {"$exists": True}}):
        try:
            ngay = doc["ngay_nop"]
            if isinstance(ngay, str):
                ngay = datetime.fromisoformat(ngay)
            month_key = ngay.strftime("%Y-%m")
            month_counts[month_key] += 1
        except:
            continue

    stats["by_month"] = dict(sorted(month_counts.items()))  # Sắp xếp theo tháng

    return stats

@router.get("/summary")
def get_candidate_summary():
    

    today = datetime.now()
    this_month = datetime(today.year, today.month, 1)

    new_candidates = collection.count_documents({"ngay_nop": {"$gte": this_month}})
    interviewed = collection.count_documents({"trang_thai": {"$regex": "phỏng vấn", "$options": "i"}})
    results_sent = collection.count_documents({"trang_thai_gui_email": "Đã gửi"})

    return {
        "new_candidates": new_candidates,
        "interviewed": interviewed,
        "results_sent": results_sent
    }


@router.get("/lookup")
def lookup_candidate(phone: str = Query(...)):
    doc = collection.find_one({"so_dien_thoai": phone})
    if not doc:
        return {"found": False}
    return {
    "found": True,
    "ho_ten": doc.get("ho_ten", ""),
    "email": doc.get("email", ""),
    "so_dien_thoai": doc.get("so_dien_thoai", ""),
    "cccd": doc.get("cccd", ""),
    "ket_qua": doc.get("ket_qua", ""),
    "trang_thai": doc.get("trang_thai", "")
}

@router.delete("/{id}")
def delete_candidate(id: str):
    try:
        object_id = ObjectId(id)
    except errors.InvalidId:
        raise HTTPException(status_code=400, detail="ID không hợp lệ")

    result = collection.delete_one({"_id": object_id})
    if result.deleted_count == 1:
        return {"success": True, "message": "Đã xóa ứng viên thành công."}
    else:
        raise HTTPException(status_code=404, detail="Không tìm thấy ứng viên để xóa")




@router.get("/{id}")
def get_candidate_by_id(id: str):
    try:
        object_id = ObjectId(id)
        doc = collection.find_one({"_id": object_id})
        if not doc:
            raise HTTPException(status_code=404, detail="Không tìm thấy ứng viên")
        doc["id"] = str(doc["_id"])
        del doc["_id"]
        return doc
    except errors.InvalidId:
        raise HTTPException(status_code=400, detail="ID không hợp lệ")
    except Exception as e:
        print("❌ Lỗi khi lấy ứng viên theo ID:", e)
        raise HTTPException(status_code=500, detail="Lỗi server")


@router.patch("/{id}/update-result")
def update_candidate_result(id: str, result: str = Query(...)):
    try:
        object_id = ObjectId(id)
    except errors.InvalidId:
        print("❌ ID không hợp lệ:", id)
        return {"error": "❌ Không thể cập nhật. ID ứng viên không hợp lệ."}

    candidate = collection.find_one({"_id": object_id})
    if not candidate:
        return {"error": "❌ Không tìm thấy ứng viên."}

    collection.update_one({"_id": object_id}, {"$set": {"ket_qua": result}})
    
    try:
        send_result_email(candidate["email"], candidate.get("ho_ten", ""), result)
    except Exception as e:
        print("❌ Lỗi gửi email:", e)

    candidate["ket_qua"] = result
    candidate["id"] = str(candidate["_id"])
    del candidate["_id"]
    return candidate

@router.put("/interviews/{interview_id}")
def update_interview(interview_id: str, data: dict = Body(...)):
    try:
        print(f"🔄 Cập nhật lịch phỏng vấn cho ID: {interview_id}")
        print(f"📦 Dữ liệu nhận được: {data}")
        
        # Chuyển đổi ID thành ObjectId
        try:
            object_id = ObjectId(interview_id)
        except errors.InvalidId:
            print(f"❌ ID không hợp lệ: {interview_id}")
            raise HTTPException(status_code=400, detail="ID không hợp lệ")
        
        # Tìm ứng viên
        candidate = collection.find_one({"_id": object_id})
        if not candidate:
            print(f"❌ Không tìm thấy ứng viên với ID: {interview_id}")
            raise HTTPException(status_code=404, detail="Không tìm thấy ứng viên")
        
        # Cập nhật thông tin lịch phỏng vấn
        update_result = collection.update_one(
            {"_id": object_id},
            {"$set": {
                "lich_phong_van.ngay": data.get("date", ""),
                "lich_phong_van.gio": data.get("time", ""),
                "lich_phong_van.dia_diem": data.get("room", ""),
                "ho_ten": data.get("candidate", candidate.get("ho_ten", ""))
            }}
        )
        
        if update_result.modified_count == 0:
            print("⚠️ Không có thay đổi nào được áp dụng")
            # Vẫn trả về thành công nếu không có thay đổi
        
        # Lấy dữ liệu đã cập nhật
        updated_candidate = collection.find_one({"_id": object_id})
        
        # Trả về dữ liệu đã được cập nhật theo định dạng cần thiết
        return {
            "id": str(updated_candidate["_id"]),
            "candidate": updated_candidate.get("ho_ten", ""),
            "date": updated_candidate.get("lich_phong_van", {}).get("ngay", ""),
            "time": updated_candidate.get("lich_phong_van", {}).get("gio", ""),
            "room": updated_candidate.get("lich_phong_van", {}).get("dia_diem", "")
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Lỗi khi cập nhật lịch phỏng vấn: {str(e)}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Lỗi server: {str(e)}")


